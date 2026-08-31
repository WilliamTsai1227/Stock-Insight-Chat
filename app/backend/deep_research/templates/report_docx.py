"""
研究報告的 Word 產出（.docx）。

吃的是 HTML 樣板那個同一份 `ReportDoc`，不是轉檔 —— 所以不會有 HTML→Word
轉換那種版面走樣，也不需要 LibreOffice 之類的重量級外部程序。

排版取向與 HTML 版一致（封面、摘要、關鍵數字帶、編號小節、側欄註記、
參考來源），但改用 Word 自己的語彙：真正的 Heading 樣式（大綱窗格與目錄可用）、
真正的超連結、真正的表格。使用者拿到手就能直接改。
"""

from __future__ import annotations

from datetime import datetime
from io import BytesIO
from typing import Any, Iterable, List, Sequence, Tuple

from .office import Block, Span, hex_rgb, office_font, paper_palette, parse_markdown
from .themes import Theme

# 行內碼用的等寬字；Windows / macOS Office 都找得到
_MONO = "Consolas"


# ─────────────────────────────────────────────────────────────
# python-docx 沒有包裝、必須直接操作 XML 的部分
# ─────────────────────────────────────────────────────────────
def _rgb(color: str):
    from docx.shared import RGBColor

    return RGBColor(*hex_rgb(color))


def _set_font(target: Any, latin: str, east_asian: str) -> None:
    """
    同時設定 latin 與 eastAsia 字型。

    python-docx 的 `font.name` 只寫 w:ascii / w:hAnsi；中文字元走的是
    w:eastAsia，沒設的話 Word 會自己挑一個字型來湊，中英文混排會忽粗忽細。
    `target` 可以是 run 或 style（兩者都有 `.font` 與可取到 rPr 的元素）。
    """
    from docx.oxml.ns import qn

    target.font.name = latin
    # run 是 `_element`、style 是 `element`；lxml 元素不能用 or 做真值判斷
    element = getattr(target, "_element", None)
    if element is None:
        element = target.element
    rpr = element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asian)


def _add_hyperlink(paragraph, text: str, url: str, color: str, fonts: Tuple[str, str]):
    """
    真正的 Word 超連結（點了會開瀏覽器）。

    python-docx 沒有這個 API：先照常建一個 run，再把它搬進 w:hyperlink 元素裡。
    """
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run(text)
    run.font.color.rgb = _rgb(color)
    run.font.underline = True
    _set_font(run, *fonts)

    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True))
    run._element.addprevious(link)
    link.append(run._element)
    return run


def _shade(cell, color: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color.lstrip("#"))
    cell._tc.get_or_add_tcPr().append(shd)


def _table_borders(
    table, *, color: str, size: int = 4, edges: Sequence[str] = ("insideH",)
) -> None:
    """只留下指定的邊；其餘明確設成 none（Word 預設樣式會自己補線）。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        if edge in edges:
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), str(size))
            element.set(qn("w:color"), color.lstrip("#"))
        else:
            element.set(qn("w:val"), "none")
            element.set(qn("w:sz"), "0")
        element.set(qn("w:space"), "0")
        borders.append(element)
    table._tbl.tblPr.append(borders)


def _cell_left_bar(cell, color: str, size: int = 18) -> None:
    """側欄註記左邊那條粗色帶（對應 HTML 版的 border-left）。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:color"), color.lstrip("#"))
    left.set(qn("w:space"), "0")
    borders.append(left)
    cell._tc.get_or_add_tcPr().append(borders)


def _page_number_field(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"                       # Word 開檔時會自行更新成實際頁碼
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


# ─────────────────────────────────────────────────────────────
# 內容組裝
# ─────────────────────────────────────────────────────────────
class _Builder:
    """把 `ReportDoc` 一段一段寫進 Word 文件。"""

    def __init__(self, document, theme: Theme) -> None:
        from docx.shared import Pt

        self.document = document
        self.theme = theme
        self.palette = paper_palette(theme)
        self.display = office_font(theme, "display")
        self.body = office_font(theme, "body")
        self.Pt = Pt

    # ── 基礎 ────────────────────────────────────────────────
    def paragraph(
        self,
        *,
        style: str | None = None,
        space_before: float = 0,
        space_after: float = 6,
        indent: float = 0,
        hanging: float = 0,
    ):
        from docx.shared import Cm

        para = self.document.add_paragraph(style=style)
        fmt = para.paragraph_format
        fmt.space_before = self.Pt(space_before)
        fmt.space_after = self.Pt(space_after)
        if indent:
            fmt.left_indent = Cm(indent)
        if hanging:
            fmt.first_line_indent = Cm(-hanging)
        return para

    def run(
        self,
        para,
        text: str,
        *,
        size: float = 10.5,
        bold: bool = False,
        color: str | None = None,
        fonts: Tuple[str, str] | None = None,
        italic: bool = False,
    ):
        run = para.add_run(text)
        run.font.size = self.Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = _rgb(color or self.palette["body"])
        _set_font(run, *(fonts or self.body))
        return run

    def spans(self, para, spans: Iterable[Span], *, size: float = 10.5) -> None:
        """把 `parse_inline()` 的結果寫進段落，粗體／行內碼／連結各自套格式。"""
        for span in spans:
            if span.link:
                _add_hyperlink(para, span.text, span.link, self.palette["accent"], self.body)
            elif span.code:
                self.run(
                    para,
                    span.text,
                    size=size - 0.5,
                    color=self.palette["ink2"],
                    fonts=(_MONO, self.body[1]),
                )
            else:
                self.run(
                    para,
                    span.text,
                    size=size,
                    bold=span.bold,
                    color=self.palette["ink"] if span.bold else self.palette["body"],
                )

    # ── 區塊 ────────────────────────────────────────────────
    def marker_item(self, marker: str, spans: Iterable[Span], *, indent: float = 0.75) -> None:
        """
        清單項目：手工畫記號 + 懸掛縮排。

        不用 Word 內建的 List Bullet / List Number 樣式 —— 那些會跨小節延續編號，
        報告裡每一節重新開始的清單會變成 7. 8. 9.，而且不同 Word 版本的
        樣式定義還不一樣。自己畫記號的結果是決定性的。
        """
        para = self.paragraph(space_after=4, indent=indent, hanging=indent)
        self.run(para, marker, bold=True, color=self.palette["accent"])
        self.spans(para, spans)

    def heading(self, text: str, *, level: int = 1) -> None:
        para = self.document.add_paragraph(style=f"Heading {min(level, 4)}")
        fmt = para.paragraph_format
        fmt.space_before = self.Pt(20 if level == 1 else 12)
        fmt.space_after = self.Pt(6)
        fmt.keep_with_next = True
        size = {1: 17, 2: 13.5, 3: 12}.get(level, 11)
        self.run(
            para,
            text,
            size=size,
            bold=True,
            color=self.palette["ink"] if level == 1 else self.palette["ink2"],
            fonts=self.display,
        )

    def markdown(self, text: str) -> None:
        """小節內文：把 Markdown 拆成區塊後逐一排版。"""
        for block in parse_markdown(text):
            if block.kind == "table":
                self.table(block)
            elif block.kind == "heading":
                self.heading("".join(s.text for s in block.spans), level=block.level + 1)
            elif block.kind == "bullet":
                self.marker_item("•　", block.spans, indent=0.75 + 0.5 * block.level)
            elif block.kind == "ordered":
                self.marker_item("–　", block.spans, indent=0.75 + 0.5 * block.level)
            else:
                para = self.paragraph(space_after=7)
                para.paragraph_format.line_spacing = 1.35
                self.spans(para, block.spans)

    def table(self, block: Block) -> None:
        rows = block.rows
        if not rows:
            return

        table = self.document.add_table(rows=len(rows), cols=len(rows[0]))
        table.autofit = True
        _table_borders(
            table,
            color=self.palette["line"],
            edges=("top", "bottom", "insideH"),
        )

        for r, row in enumerate(rows):
            for c, cell_spans in enumerate(row):
                cell = table.cell(r, c)
                para = cell.paragraphs[0]
                para.paragraph_format.space_before = self.Pt(3)
                para.paragraph_format.space_after = self.Pt(3)
                if r == 0:
                    _shade(cell, self.palette["bg-soft"])
                    for span in cell_spans:
                        self.run(
                            para,
                            span.text,
                            size=9.5,
                            bold=True,
                            color=self.palette["ink2"],
                        )
                else:
                    self.spans(para, cell_spans, size=9.5)

        self.paragraph(space_after=4)      # 表格後補一個空段，不然下一段會黏上來

    def callout(self, text: str) -> None:
        table = self.document.add_table(rows=1, cols=1)
        _table_borders(table, color=self.palette["accent-line"], edges=())
        cell = table.cell(0, 0)
        _shade(cell, self.palette["accent-soft"])
        _cell_left_bar(cell, self.palette["accent"])

        label = cell.paragraphs[0]
        label.paragraph_format.space_before = self.Pt(6)
        label.paragraph_format.space_after = self.Pt(2)
        self.run(label, "重點", size=8.5, bold=True, color=self.palette["accent"])

        body = cell.add_paragraph()
        body.paragraph_format.space_after = self.Pt(6)
        self.run(body, text, size=10.5, color=self.palette["ink2"])
        self.paragraph(space_after=4)


def _setup_document(document, builder: _Builder) -> None:
    """頁面尺寸、預設樣式、頁尾。"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    section = document.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    normal = document.styles["Normal"]
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = _rgb(builder.palette["body"])
    _set_font(normal, *builder.body)
    normal.paragraph_format.space_after = Pt(6)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    builder.run(footer, "", size=9, color=builder.palette["muted"])
    _page_number_field(footer)


def render_report_docx(doc: Any, *, model: str, query: str, theme: Theme) -> bytes:
    """把 `ReportDoc` 轉成 .docx 的位元組內容。"""
    from docx import Document

    document = Document()
    builder = _Builder(document, theme)
    _setup_document(document, builder)
    palette = builder.palette

    # ── 封面區 ──────────────────────────────────────────────
    kicker = builder.paragraph(space_after=4)
    builder.run(kicker, "深度研究報告", size=9, bold=True, color=palette["accent"])

    title = builder.paragraph(space_after=6)
    builder.run(
        title, doc.title, size=26, bold=True, color=palette["ink"], fonts=builder.display
    )

    if doc.subtitle:
        subtitle = builder.paragraph(space_after=12)
        builder.run(subtitle, doc.subtitle, size=13, color=palette["ink2"])

    meta = builder.paragraph(space_after=2)
    builder.run(
        meta,
        f"研究題目：{query}",
        size=9,
        color=palette["muted"],
    )
    meta2 = builder.paragraph(space_after=14)
    builder.run(
        meta2,
        f"模型：{model}　·　產出時間：{datetime.now().strftime('%Y-%m-%d %H:%M')}",
        size=9,
        color=palette["muted"],
    )

    # ── 摘要 ────────────────────────────────────────────────
    if doc.executive_summary:
        builder.heading("摘要")
        for n, item in enumerate(doc.executive_summary, start=1):
            builder.marker_item(f"{n:02d}　", [Span(item)])

    # ── 關鍵數字 ────────────────────────────────────────────
    if doc.key_figures:
        builder.heading("關鍵數字")
        table = document.add_table(rows=len(doc.key_figures), cols=2)
        _table_borders(table, color=palette["line"], edges=("top", "bottom", "insideH"))
        for row_index, figure in enumerate(doc.key_figures):
            value_cell = table.cell(row_index, 0)
            value_para = value_cell.paragraphs[0]
            value_para.paragraph_format.space_before = builder.Pt(6)
            value_para.paragraph_format.space_after = builder.Pt(6)
            builder.run(
                value_para,
                figure.value,
                size=15,
                bold=True,
                color=palette["accent"],
                fonts=builder.display,
            )
            label_para = table.cell(row_index, 1).paragraphs[0]
            label_para.paragraph_format.space_before = builder.Pt(6)
            label_para.paragraph_format.space_after = builder.Pt(6)
            builder.run(label_para, figure.label, size=10, color=palette["body"])
        builder.paragraph(space_after=4)

    # ── 正文小節 ────────────────────────────────────────────
    for n, section in enumerate(doc.sections, start=1):
        builder.heading(f"{n:02d}　{section.heading}")
        builder.markdown(section.body_markdown)
        if section.callout.strip():
            builder.callout(section.callout)

    # ── 尚待釐清 ────────────────────────────────────────────
    if doc.open_questions:
        builder.heading("尚待釐清")
        for question in doc.open_questions:
            builder.marker_item("？　", [Span(question)])

    # ── 參考來源 ────────────────────────────────────────────
    if doc.references:
        builder.heading("參考來源")
        for n, ref in enumerate(doc.references, start=1):
            para = builder.paragraph(space_after=4, indent=0.9, hanging=0.9)
            builder.run(para, f"{n:02d}　", bold=True, color=palette["accent"])
            # 使用者上傳的檔案沒有網址，這時只寫標題，不要生一個點不開的連結
            if ref.url.startswith(("http://", "https://")):
                _add_hyperlink(para, ref.title or ref.url, ref.url, palette["accent"], builder.body)
                builder.run(para, f"　{ref.url}", size=8.5, color=palette["muted"])
            else:
                builder.run(para, ref.title or ref.url, color=palette["ink2"])

    # ── 頁尾聲明 ────────────────────────────────────────────
    disclaimer = builder.paragraph(space_before=18, space_after=0)
    builder.run(
        disclaimer,
        f"本報告由 Insight 深度研究以 {model} 產生，內容可能有誤，請於引用前自行查證來源。",
        size=8.5,
        color=palette["muted"],
        italic=True,
    )

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
